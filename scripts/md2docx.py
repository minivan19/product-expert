#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

"""
Markdown转Word文档转换器
支持基本Markdown语法转换为Word文档
"""

import argparse
import os
import sys
import re
from pathlib import Path
from typing import Optional, List, Tuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class MarkdownToDocxConverter:
    """Markdown转Word转换器"""
    
    def __init__(self, template_path: Optional[str] = None, debug: bool = False):
        self.debug = debug
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            self.doc = Document()
        
        self.current_paragraph = None
        self.current_table = None
        
    def convert(self, markdown_text: str, image_dir: str = None):
        """转换Markdown文本为Word文档"""
        lines = markdown_text.split('\n')
        
        in_code_block = False
        code_content = []
        
        for line in lines:
            # 跳过空行
            if not line.strip():
                in_code_block = False
                continue
            
            # 代码块处理
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_content = []
                else:
                    # 代码块结束
                    self._add_code_block('\n'.join(code_content))
                    in_code_block = False
                    code_content = []
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # 处理表格分隔行（跳过|---|---|这样的行）
            if self._is_table_separator(line):
                continue
            
            # 标题处理
            if line.startswith('# '):
                self.doc.add_heading(line[2:], 0)
            elif line.startswith('## '):
                self.doc.add_heading(line[3:], 1)
            elif line.startswith('### '):
                self.doc.add_heading(line[4:], 2)
            elif line.startswith('#### '):
                self.doc.add_heading(line[5:], 3)
            elif line.startswith('##### '):
                self.doc.add_heading(line[6:], 4)
            elif line.startswith('###### '):
                self.doc.add_heading(line[7:], 5)
            # 表格处理
            elif line.startswith('|'):
                self._process_table_row(line)
            # 列表处理
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                self._add_paragraph_with_formatting(line.strip()[2:], style='List Bullet')
            elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ', '0. ')):
                # 有序列表
                self._add_paragraph_with_formatting(line.strip(), style='List Number')
            # 引用处理
            elif line.strip().startswith('> '):
                self.doc.add_paragraph(line.strip()[2:], style='Quote')
            # 水平线
            elif line.strip() == '---':
                self.doc.add_paragraph()
            # 普通段落
            else:
                self._add_paragraph_with_formatting(line)
    
    def _is_table_separator(self, line: str) -> bool:
        """检查是否是表格分隔行"""
        stripped = line.strip()
        if not stripped.startswith('|'):
            return False
        # 检查是否只包含|、-和空格
        content = stripped.replace('|', '').replace('-', '').replace(' ', '')
        return len(content) == 0
    
    def _process_table_row(self, line: str):
        """处理表格行"""
        # 解析表格行
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        
        if len(cells) == 0:
            return
        
        # 检查是否是表头
        is_header = True
        if self.current_table:
            if len(self.current_table.columns) != len(cells):
                # 列数不同，创建新表格
                self.current_table = None
                is_header = True
            else:
                # 列数相同，检查表头内容
                prev_header = [cell.text.strip() for cell in self.current_table.rows[0].cells]
                if prev_header == cells:
                    # 表头完全相同，创建新表格
                    self.current_table = None
                    is_header = True
                else:
                    # 表头内容不同（即使列数相同），也创建新表格
                    self.current_table = None
                    is_header = True
        
        if is_header:
            # 创建新表格
            self.current_table = self.doc.add_table(rows=1, cols=len(cells))
            self.current_table.style = 'Table Grid'

            # 添加表头
            header_cells = self.current_table.rows[0].cells
            for i, cell_text in enumerate(cells):
                if i < len(header_cells):
                    # 去掉星号再写入
                    header_cells[i].text = cell_text.replace('**', '')
        else:
            # 添加数据行
            if self.current_table:
                row = self.current_table.add_row()
                for i, cell_text in enumerate(cells):
                    if i < len(row.cells):
                        # 去掉星号再写入
                        row.cells[i].text = cell_text.replace('**', '')
    
    def _add_paragraph_with_formatting(self, text: str, style: str = None):
        """添加段落并处理基本格式"""
        paragraph = self.doc.add_paragraph(style=style) if style else self.doc.add_paragraph()

        # 用 findall 精确提取 bold 内容，避免正则 split 的歧义问题
        bold_pattern = re.compile(r'\*\*([^*]+)\*\*')
        last_end = 0
        for match in bold_pattern.finditer(text):
            # 添加匹配之前的普通文本
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            # 添加 bold 文本
            run = paragraph.add_run(match.group(1))
            run.bold = True
            last_end = match.end()
        # 添加剩余文本
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
    
    def _add_code_block(self, code: str):
        """添加代码块"""
        paragraph = self.doc.add_paragraph(code)
        # 尝试设置Code样式，如果不存在则忽略
        try:
            paragraph.style = 'Code'
        except:
            pass  # 样式不存在时忽略
    
    def save(self, output_path: str):
        """Saving document"""
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)


def convert_markdown_to_docx(input_file: str, output_file: str, 
                             template: Optional[str] = None,
                             image_dir: Optional[str] = None,
                             debug: bool = False) -> bool:
    """转换Markdown文件为Word文档"""
    try:
        if not os.path.exists(input_file):
            print(f"ERROR Input file not found: {input_file}")
            return False
        
        # Reading Markdown file
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        
        if debug:
            print(f"FILE Reading Markdown file: {input_file} ({len(markdown_text)} 字符)")
        
        # Creating converter
        converter = MarkdownToDocxConverter(template, debug)
        
        # 转换
        converter.convert(markdown_text, image_dir)
        
        # 保存
        converter.save(output_file)
        
        print(f"OK Conversion successful: {input_file} → {output_file}")
        return True
        
    except Exception as e:
        print(f"ERROR Conversion failed: {e}")
        if debug:
            import traceback
            traceback.print_exc()
        return False


def list_available_styles(template_path: Optional[str] = None):
    """列出可用样式"""
    try:
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
            print(f"📋 模板中的样式: {template_path}")
        else:
            doc = Document()
            print("📋 默认样式:")
        
        for style in doc.styles:
            print(f"  - {style.name} ({style.type})")
    
    except Exception as e:
        print(f"ERROR 列出样式失败: {e}")


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='Markdown转Word文档转换器')
    parser.add_argument('--input', '-i', required=True, help='输入Markdown文件路径')
    parser.add_argument('--output', '-o', required=True, help='输出Word文件路径')
    parser.add_argument('--template', '-t', help='Word模板文件路径')
    parser.add_argument('--image-dir', '-d', help='图片目录路径')
    parser.add_argument('--list-styles', action='store_true', help='列出可用样式')
    parser.add_argument('--debug', action='store_true', help='启用调试模式')
    parser.add_argument('--encoding', default='utf-8', help='文件编码（默认: utf-8）')
    
    args = parser.parse_args()
    
    if args.list_styles:
        list_available_styles(args.template)
        return
    
    # 确保输出目录存在
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    # 执行转换
    success = convert_markdown_to_docx(
        input_file=args.input,
        output_file=args.output,
        template=args.template,
        image_dir=args.image_dir,
        debug=args.debug
    )
    
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
